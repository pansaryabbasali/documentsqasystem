"""DOCX extraction via python-docx, one TextBlock per heading-delimited section.

Word documents have no fixed pages (pagination is a rendering artifact), so
the honest citation unit is the SECTION: text is grouped under the most
recent heading, and the locator names it — e.g. ``section "Carry-Over
Policy"`` — which is more useful to a reader than a page number anyway.

Headings are detected two ways: real ``Heading*`` styles, OR the pattern
found in this corpus (and most real offices): a short, fully-bold
``Normal`` paragraph. Tables are emitted as their own blocks (``table N``)
because python-docx exposes them separately from the paragraph stream.
"""

from __future__ import annotations

from collections.abc import Iterator
from pathlib import Path

import docx

from .base import TextBlock


class DocxLoader:
    suffixes: tuple[str, ...] = (".docx",)

    def load(self, path: Path) -> Iterator[TextBlock]:
        document = docx.Document(str(path))
        section: str | None = None
        lines: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if self._is_heading(paragraph, text):
                yield from self._flush(lines, path.name, section)
                section, lines = text, [text]  # heading included in its own section's text
            else:
                lines.append(text)
        yield from self._flush(lines, path.name, section)

        for number, table in enumerate(document.tables, start=1):
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
            ]
            body = "\n".join(row for row in rows if row.strip(" |"))
            if body:
                yield TextBlock(text=body, source=path.name, locator=f"table {number}")

    @staticmethod
    def _is_heading(paragraph: docx.text.paragraph.Paragraph, text: str) -> bool:
        if paragraph.style.name.startswith("Heading"):
            return True
        runs = [run for run in paragraph.runs if run.text.strip()]
        return bool(runs) and all(run.bold for run in runs) and len(text) <= 60

    @staticmethod
    def _flush(lines: list[str], source: str, section: str | None) -> Iterator[TextBlock]:
        if lines:
            locator = f'section "{section}"' if section else "preamble"
            yield TextBlock(text="\n".join(lines), source=source, locator=locator)
        lines.clear()
